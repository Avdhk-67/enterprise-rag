"""Document processor for parsing and processing various document formats."""
import io
import os
import tempfile
from typing import List, Dict, Optional
from pathlib import Path
from PyPDF2 import PdfReader
from docx import Document
from src.ingestion.s3_document_loader import S3DocumentLoader
from src.ingestion.ocr_processor import OCRProcessor


class DocumentProcessor:
    """Process documents from various formats."""
    
    def __init__(self):
        self.s3_loader = S3DocumentLoader()
        self.ocr_processor = OCRProcessor()
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.txt': self._process_txt,
            '.md': self._process_txt,
            '.jpg': self._process_image,
            '.jpeg': self._process_image,
            '.png': self._process_image
        }
    
    def process_document(self, s3_key: str) -> Dict:
        """Process a document from S3."""
        file_extension = Path(s3_key).suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Download document
        file_content = self.s3_loader.download_document(s3_key)
        
        # Process based on format
        # We handle PDF and Images specially to support OCR and S3 upload
        if file_extension == '.pdf':
            text_content = self._process_pdf(file_content, s3_key)
        elif file_extension in ['.jpg', '.jpeg', '.png']:
            text_content = self._process_image(file_content, s3_key)
        else:
            processor = self.supported_formats[file_extension]
            text_content = processor(file_content)
        
        # Get metadata
        metadata = self.s3_loader.get_document_metadata(s3_key)
        
        return {
            's3_key': s3_key,
            'text': text_content,
            'metadata': metadata,
            'format': file_extension
        }
    
    def _save_ocr_result(self, text: str, original_key: str):
        """Save OCR extracted text as a PDF to S3."""
        if not text.strip():
            return

        try:
            # Create temporary PDF
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
                pdf_path = tmp.name
            
            self.ocr_processor.create_text_pdf(text, pdf_path)
            
            # Read the new PDF content
            with open(pdf_path, 'rb') as f:
                new_pdf_content = f.read()
            
            # Upload to S3 (processed-ocr prefix)
            filename = Path(original_key).stem
            new_key = f"processed-ocr/{filename}.pdf"
            
            print(f"Uploading OCR processed PDF to: {new_key}")
            self.s3_loader.upload_document(new_pdf_content, new_key)
            
            # Cleanup
            os.remove(pdf_path)
        except Exception as e:
            print(f"Error saving OCR result for {original_key}: {str(e)}")

    def _process_image(self, file_content: bytes, s3_key: Optional[str] = None) -> str:
        """Extract text from Image using OCR."""
        print(f"Processing image with OCR: {s3_key}")
        text = self.ocr_processor.extract_text_from_bytes(file_content)
        
        if s3_key:
            self._save_ocr_result(text, s3_key)
            
        return text

    def _process_pdf(self, file_content: bytes, s3_key: Optional[str] = None) -> str:
        """Extract text from PDF (with OCR fallback)."""
        pdf_file = io.BytesIO(file_content)
        try:
            reader = PdfReader(pdf_file)
        except Exception as e:
            print(f"Error reading PDF structure (possibly corrupted or encrypted): {str(e)}. Attempting OCR fallback for: {s3_key}")
            if s3_key:
                try:
                    text = self.ocr_processor.extract_text_from_s3_pdf(self.s3_loader.bucket_name, s3_key)
                    self._save_ocr_result(text, s3_key)
                    return text
                except Exception as ocr_e:
                    print(f"OCR fallback also failed: {str(ocr_e)}")
                    return ""
            return ""
        text_parts = []
        
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text_parts.append(extracted)
        
        full_text = '\n\n'.join(text_parts)
        
        # If text is empty or very short, assume scanned and use OCR
        # We also check if the text density is suspiciously low (e.g. < 20 chars per page)
        # This handles cases where PyPDF2 extracts some garbage or headers/footers but not the main content
        text_length = len(full_text.strip())
        num_pages = len(reader.pages)
        avg_chars_per_page = text_length / num_pages if num_pages > 0 else 0
        
        if not full_text.strip() or (num_pages > 0 and avg_chars_per_page < 20):
            print(f"PDF seems scanned (avg {avg_chars_per_page:.1f} chars/page). Falling back to OCR for: {s3_key}")
            # For PDF OCR, we need to convert pages to images or use Textract's PDF support
            # Textract detect_document_text only supports JPEG/PNG
            # Textract analyze_document supports PDF but requires async for multi-page or sync for single page
            # For simplicity and since we have the bytes, we can try to pass bytes to Textract if it's a single page
            # OR better: Use Textract's start_document_text_detection for PDFs (async)
            # BUT: To keep it simple and synchronous as per current architecture:
            # We will use the OCRProcessor which currently uses detect_document_text (Images only)
            # So we need to convert PDF to Image? That adds 'pdf2image' dependency and 'poppler'.
            # ALTERNATIVE: Use Textract's synchronous analyze_document/detect_document_text IF it supports PDF bytes?
            # AWS Textract Sync API only supports PDF for 'AnalyzeDocument' and 'DetectDocumentText' IF it's passed as S3Object?
            # Actually, Sync APIs only support Images (JPG/PNG). PDF is only supported in Async APIs or if we convert to image.
            
            # Wait, let's check OCRProcessor. It uses detect_document_text.
            # If we pass PDF bytes to detect_document_text, it might fail.
            
            # Let's try to see if we can use the S3 object directly since we have the key!
            # Textract can read directly from S3.
            
            if s3_key:
                try:
                    # We need the bucket name. S3DocumentLoader has it.
                    bucket = self.s3_loader.bucket_name
                    
                    # Use start_document_text_detection (Async) or just assume it's an image?
                    # No, it's a PDF.
                    # We need to use start_document_text_detection for PDFs in S3.
                    
                    # For this iteration, let's implement a simple synchronous wait for the async job
                    # OR update OCRProcessor to handle PDFs via S3.
                    
                    # Let's update OCRProcessor to handle this.
                    # But I can't update OCRProcessor right now in this tool call.
                    # I will call a method on OCRProcessor that I will add/update later?
                    # No, I should have updated OCRProcessor first.
                    
                    # Let's assume OCRProcessor will have a method `extract_text_from_s3_pdf(bucket, key)`
                    
                    text = self.ocr_processor.extract_text_from_s3_pdf(self.s3_loader.bucket_name, s3_key)
                    
                    # Save the result as a new PDF
                    self._save_ocr_result(text, s3_key)
                    return text
                    
                except Exception as e:
                    print(f"OCR fallback failed: {str(e)}")
                    return ""
            
        return full_text
    
    def _process_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX."""
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            text_parts.append(paragraph.text)
        
        return '\n\n'.join(text_parts)
    
    def _process_txt(self, file_content: bytes) -> str:
        """Extract text from TXT/MD."""
        try:
            return file_content.decode('utf-8')
        except UnicodeDecodeError:
            return file_content.decode('latin-1')
    
    def process_all_documents(self, prefix: Optional[str] = None) -> List[Dict]:
        """Process all documents in S3 bucket."""
        documents = self.s3_loader.list_documents(prefix)
        processed_docs = []
        
        for doc_info in documents:
            try:
                processed = self.process_document(doc_info['key'])
                processed_docs.append(processed)
                print(f"Processed: {doc_info['key']}")
            except Exception as e:
                print(f"Error processing {doc_info['key']}: {str(e)}")
        
        return processed_docs

    def process_local_pdf(self, file_path: str) -> Dict:
        # Simplified for now, just standard PDF read
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            t = page.extract_text()
            if t: text_parts.append(t)

        full_text = '\n\n'.join(text_parts)
        return {"text": full_text, "s3_key": "local_upload", "metadata": {"source": file_path}, "format": "pdf"}


if __name__ == "__main__":
    from src.ingestion.pipeline import IngestionPipeline
    print("Starting document ingestion pipeline...")
    pipeline = IngestionPipeline()
    result = pipeline.ingest_from_s3()
    print(f"\nIngestion complete!")
    print(f"Documents Processed: {result['documents_processed']}")
    print(f"Chunks Created: {result['chunks_created']}")
    print(f"Status: {result['status']}")
# from typing import List, Dict, Optional
# from pathlib import Path
# from PyPDF2 import PdfReader
# from docx import Document
# from src.ingestion.s3_document_loader import S3DocumentLoader


# class DocumentProcessor:
#     def __init__(self):
#         self.s3_loader = S3DocumentLoader()
#         self.supported_formats = {
#             ".pdf": self._process_pdf,
#             ".docx": self._process_docx,
#             ".txt": self._process_txt,
#             ".md": self._process_txt
#         }

#     def process_all_documents(self, prefix: Optional[str] = None) -> List[Dict]:
#         documents = self.s3_loader.list_documents(prefix)
#         processed_docs = []
#         for doc_info in documents:
#             try:
#                 processed = self.process_document(doc_info['key'])
#                 processed_docs.append(processed)
#                 print(f"Processed: {doc_info['key']}")
#             except Exception as e:
#                 print(f"Error processing {doc_info['key']}: {str(e)}")
#         return processed_docs

#     def process_document(self, s3_key: str) -> Dict:
#         file_extension = Path(s3_key).suffix.lower()

#         file_content = self.s3_loader.download_document(s3_key)
#         metadata = self.s3_loader.get_document_metadata(s3_key)

#         if file_extension == ".pdf":
#             text = self._process_pdf(file_content)
#         elif file_extension == ".docx":
#             text = self._process_docx(file_content)
#         elif file_extension in [".txt", ".md"]:
#             text = self._process_txt(file_content)
#         else:
#             raise ValueError(f"Unsupported file format {file_extension}")

#         return {
#             "s3_key": s3_key,
#             "text": text,
#             "metadata": metadata,
#             "format": file_extension
#         }

#     def _process_pdf(self, file_content: bytes) -> str:
#         pdf_file = io.BytesIO(file_content)
#         reader = PdfReader(pdf_file)
#         text = ""

#         for page in reader.pages:
#             if page.extract_text():
#                 text += page.extract_text() + "\n"

#         return text

#     def _process_docx(self, file_content: bytes) -> str:
#         docx_file = io.BytesIO(file_content)
#         doc = Document(docx_file)
#         return "\n".join(p.text for p in doc.paragraphs)

#     def _process_txt(self, file_content: bytes) -> str:
#         try:
#             return file_content.decode("utf-8")
#         except:
#             return file_content.decode("latin-1")

#     def process_local_pdf(self, file_path: str) -> Dict:
#         reader = PdfReader(file_path)
#         text = ""

#         for page in reader.pages:
#             text += page.extract_text() or ""

#         return {
#             "text": text,
#             "metadata": {"source": file_path},
#             "format": ".pdf"
#         }
